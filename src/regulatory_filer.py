"""
src/regulatory_filer.py — Molecular Manufacturability & Risk Assessment Report Generator
=============================================================================
ProtePilot — Milestone 24 · Version 2.0

Auto-generates a Molecular Manufacturability & Risk Assessment Report with
comprehensive, paragraph-length scientific narratives for every section.

Upgrade from M22 v1.0 (bullet points) → M24 v2.0 (deep narrative):
  - Each section now generates multi-paragraph scientific justification
  - Context-aware narrative engine consumes raw metrics and produces
    regulatory-grade prose with scientific rationale
  - Optional LLM enhancement via OpenAI API (falls back to template engine)
  - ADA/immunogenicity and scale-up data integrated from M23

Manufacturability & Risk Assessment Structure (ICH M4Q reference)
----------------------------------
  3.2.S  Drug Substance
    3.2.S.1    General Information (Structure, Properties)
    3.2.S.2    Manufacture (Upstream + Downstream Process)
    3.2.S.3    Characterisation (Structure Elucidation, Impurities)
    3.2.S.4    Control of Drug Substance
    3.2.S.7    Stability

  3.2.P  Drug Product
    3.2.P.1    Description and Composition
    3.2.P.2    Pharmaceutical Development
    3.2.P.3    Manufacture
    3.2.P.8    Stability

Output: Markdown (for display) or Word .docx (for download)
"""

from __future__ import annotations

import io
import logging
import time
from typing import Any, Dict, List, Optional

log = logging.getLogger(__name__)


# ===========================================================================
# 1. Data Extraction (expanded for M24)
# ===========================================================================

def _get(d: Dict, *keys, default=None):
    """Nested safe get."""
    for k in keys:
        if isinstance(d, dict):
            d = d.get(k, default)
        else:
            return default
    return d


def _num(val, fmt=".2f") -> str:
    """Format a value as number string, or return '[pending]' if not numeric."""
    if val is None or val == "[pending]" or val == "[N/A]":
        return "[pending]"
    try:
        return f"{float(val):{fmt}}"
    except (ValueError, TypeError):
        return str(val)


def _extract_filing_data(session_data: Dict[str, Any]) -> Dict[str, Any]:
    """Extract all data needed for Manufacturability & Risk Assessment."""
    intent = session_data.get("last_intent") or {}
    # Fallback: if last_intent is empty/missing, pull from workspace store
    if not intent or not intent.get("sequence"):
        _ws = session_data.get("_workspace_store")
        if _ws and hasattr(_ws, "get_active"):
            _active_ws = _ws.get_active()
            if _active_ws and isinstance(_active_ws, dict):
                _ws_intent = _active_ws.get("intent")
                if _ws_intent and isinstance(_ws_intent, dict):
                    intent = _ws_intent
    up = session_data.get("upstream_result_dict") or {}
    doe = session_data.get("doe_result_dict") or {}
    cogs = session_data.get("cogs_result_dict") or {}

    # Developability
    dev_score = None
    agg_risk = None
    stability = None
    _ws = session_data.get("_workspace_store")
    _active = None
    if _ws and hasattr(_ws, "get_active"):
        _active = _ws.get_active()
    if _active and isinstance(_active, dict):
        _cache = _active.get("analysis_cache") or {}
        _dev = (_cache.get("dev_result") or {}).get("data", {})
        if _dev:
            dev_score = _dev.get("composite_score")
            preds = _dev.get("predictions", {})
            agg_risk = preds.get("agg_risk")
            stability = preds.get("stability")

    # ADA / immunogenicity (M23)
    ada = session_data.get("ada_result")
    ada_risk_level = getattr(ada, "ada_risk_level", None) if ada else None
    ada_risk_score = getattr(ada, "ada_risk_score", None) if ada else None
    ada_n_hotspots = getattr(ada, "n_high_risk", 0) if ada else 0

    # Scale-up (M23)
    su = session_data.get("scaleup_result")
    su_titer = getattr(su, "predicted_titer_large", None) if su else None
    su_strategy = getattr(su, "recommended_strategy", None) if su else None

    return {
        "name": intent.get("name", intent.get("molecule_name", "[Candidate Name]")),
        "pI": intent.get("pI", "[pI]"),
        "mw": intent.get("mw", "[MW]"),
        "sequence": intent.get("sequence", ""),
        "seq_length": len(intent.get("sequence", "")),
        "num_chains": intent.get("num_chains", 1),
        "hydrophobicity": intent.get("hydrophobicity", "[N/A]"),
        "titer": up.get("final_titer", "[pending]"),
        "peak_vcd": up.get("peak_vcd", "[pending]"),
        "viability": up.get("viability_at_harvest", "[pending]"),
        "temp_shift_day": up.get("temp_shift_day", 5),
        "seed_density": up.get("seed_density", 0.5),
        "optimal_ph": doe.get("optimal_ph", "[pending]"),
        "optimal_gradient": doe.get("optimal_gradient", "[pending]"),
        "resolution": doe.get("optimal_resolution", "[pending]"),
        "ds_yield": doe.get("optimal_yield", "[pending]"),
        "cogs_per_gram": cogs.get("cogs_per_gram", "[pending]"),
        "cost_rating": cogs.get("cost_rating", "[pending]"),
        "batch_output_g": cogs.get("batch_output_g", "[pending]"),
        "dev_score": round(dev_score, 4) if isinstance(dev_score, (int, float)) else (dev_score or "[pending]"),
        "agg_risk": round(agg_risk, 4) if isinstance(agg_risk, (int, float)) else (agg_risk or "[pending]"),
        "stability": round(stability, 4) if isinstance(stability, (int, float)) else (stability or "[pending]"),
        "form_ph": session_data.get("formulation_buffer_ph", "[pending]"),
        "form_buffer": session_data.get("formulation_buffer_type", "[pending]"),
        "excipients": session_data.get("formulation_excipients", []),
        "glycoform": session_data.get("glycoform_profile", "standard_cho"),
        "ada_risk_level": ada_risk_level,
        "ada_risk_score": ada_risk_score,
        "ada_n_hotspots": ada_n_hotspots,
        "su_titer": su_titer,
        "su_strategy": su_strategy,
    }


def _safe_ph_pi_gap(pI_val, ph_val) -> str:
    """Compute pH-pI gap string, or return placeholder if values aren't numeric."""
    try:
        gap = abs(float(pI_val) - float(ph_val))
        return f"{gap:.1f}"
    except (ValueError, TypeError):
        return "[pending]"


# ===========================================================================
# 2. Deep Narrative Engine — Scientific Prose Generators
# ===========================================================================

def _narrate_structure(d: Dict) -> str:
    """Generate 3.2.S.1 — General Information (Structure & Properties) narrative."""
    name = d["name"]
    pI = d["pI"]
    mw = d["mw"]
    seq_len = d["seq_length"]
    hydro = d["hydrophobicity"]
    glyco = d["glycoform"]
    dev = d["dev_score"]
    agg = d["agg_risk"]
    stab = d["stability"]
    seq_preview = d["sequence"][:60] + "..." if len(d["sequence"]) > 60 else d["sequence"]

    paras = []

    paras.append(
        f"{name} is a recombinant monoclonal antibody (IgG) produced in Chinese Hamster "
        f"Ovary (CHO) cells. The molecule consists of {d['num_chains']} polypeptide chain(s) "
        f"with a combined amino acid length of {seq_len} residues and a calculated molecular "
        f"weight of {mw} kDa. The theoretical isoelectric point (pI) is {pI}, which places "
        f"this molecule in the {'basic' if _is_numeric_gt(pI, 7.5) else 'near-neutral'} "
        f"range typical of IgG1 antibodies."
    )

    if _is_numeric(hydro):
        hydro_val = float(hydro)
        if hydro_val > 0.45:
            hydro_assessment = (
                f"The surface hydrophobicity index ({hydro}) is elevated, indicating a higher "
                f"propensity for non-specific interactions and potential aggregation via "
                f"hydrophobic patch exposure. This warrants careful formulation development "
                f"with appropriate surfactant (e.g., polysorbate 80) to shield exposed "
                f"hydrophobic regions."
            )
        elif hydro_val > 0.3:
            hydro_assessment = (
                f"The surface hydrophobicity index ({hydro}) is within the moderate range. "
                f"Standard formulation conditions with appropriate surfactant concentration "
                f"are expected to provide adequate colloidal stability."
            )
        else:
            hydro_assessment = (
                f"The surface hydrophobicity index ({hydro}) is low, suggesting favorable "
                f"colloidal behavior and reduced risk of non-specific binding. This is a "
                f"positive attribute for both manufacturability and pharmacokinetic clearance."
            )
        paras.append(hydro_assessment)

    if _is_numeric(dev):
        dev_val = float(dev)
        if dev_val < 0.35:
            dev_text = (
                f"The composite developability score of {dev_val:.2f} indicates significant "
                f"developability challenges. This score integrates predictions for aggregation "
                f"propensity, thermal stability, and viscosity risk. At this level, the "
                f"molecule requires engineering optimization or careful process development "
                f"to ensure commercial viability."
            )
        elif dev_val < 0.65:
            dev_text = (
                f"The composite developability score of {dev_val:.2f} places this molecule "
                f"in the moderate-risk category. While the overall profile is acceptable for "
                f"clinical development, specific liabilities identified during characterization "
                f"should be addressed through formulation optimization and process controls."
            )
        else:
            dev_text = (
                f"The composite developability score of {dev_val:.2f} indicates a favorable "
                f"developability profile. Predicted aggregation risk, thermal stability, and "
                f"solution viscosity are within acceptable ranges for commercial manufacturing "
                f"and clinical administration."
            )
        paras.append(dev_text)

    if _is_numeric(agg):
        agg_val = float(agg)
        if agg_val > 0.5:
            paras.append(
                f"In-silico structural assessment indicates a high propensity for "
                f"high-molecular-weight (HMW) species formation (predicted aggregation risk: "
                f"{agg_val:.2f}). This is primarily driven by exposed hydrophobic patches "
                f"and unfavorable charge distribution. To mitigate this liability during "
                f"downstream processing, a formulation strategy incorporating stabilizing "
                f"excipients at an optimized pH is recommended, along with SEC-HPLC monitoring "
                f"at each process step to control HMW levels below 5%."
            )
        elif agg_val > 0.3:
            paras.append(
                f"The predicted aggregation risk ({agg_val:.2f}) is moderate. Standard process "
                f"controls including appropriate hold times, temperature management, and "
                f"formulation with lyoprotectant excipients should be sufficient to maintain "
                f"HMW species below acceptance criteria."
            )

    glyco_text = {
        "standard_cho": (
            "The molecule is expressed with a standard CHO glycosylation profile, "
            "predominantly G0F and G1F glycoforms at the Fc N297 site. This profile "
            "is consistent with well-characterized IgG1 therapeutic antibodies and "
            "provides appropriate Fc-gamma receptor binding for the intended mechanism of action."
        ),
        "high_mannose": (
            "The molecule exhibits a high-mannose glycosylation profile, which may "
            "result in faster serum clearance via the mannose receptor pathway. This "
            "glycoform profile should be monitored during process development, as "
            "mannose content is sensitive to culture conditions."
        ),
        "sialylated": (
            "The molecule carries a sialylated glycosylation profile, which is associated "
            "with extended serum half-life and anti-inflammatory properties via the "
            "FcgammaRIIB pathway. Process conditions should be optimized to maintain consistent "
            "sialylation levels across manufacturing campaigns."
        ),
    }
    paras.append(glyco_text.get(d["glycoform"], glyco_text["standard_cho"]))

    # Format as Markdown
    header = f"""## 3.2.S Drug Substance

### 3.2.S.1 General Information

#### 3.2.S.1.1 Nomenclature

**INN / Working Name:** {name}
**Type:** Monoclonal Antibody (IgG)
**CAS Registry Number:** [To be assigned]

#### 3.2.S.1.2 Structure and Properties

**Primary Sequence (N-terminal preview):**
```
{seq_preview}
```

| Property | Value |
|----------|-------|
| Molecular Weight | {mw} kDa |
| Isoelectric Point (pI) | {pI} |
| Total Amino Acids | {seq_len} |
| Number of Chains | {d['num_chains']} |
| Surface Hydrophobicity | {hydro} |
| Glycoform Profile | {glyco} |
| Developability Score | {dev} |
| Aggregation Risk | {agg} |
| Thermal Stability | {stab} |

"""
    return header + "\n\n".join(paras)


def _narrate_manufacture(d: Dict) -> str:
    """Generate 3.2.S.2 — Manufacturing Process narrative."""
    titer = d["titer"]
    vcd = d["peak_vcd"]
    viab = d["viability"]
    ph = d["optimal_ph"]
    grad = d["optimal_gradient"]
    res = d["resolution"]
    ds_yield = d["ds_yield"]
    cogs = d["cogs_per_gram"]

    ds_yield_str = f"{ds_yield:.1%}" if isinstance(ds_yield, (int, float)) else str(ds_yield)

    paras = []

    # Upstream narrative
    up_para = (
        f"The drug substance is manufactured using a fed-batch CHO cell culture process. "
        f"Cells are inoculated at a seed density of {d['seed_density']} x10^6 viable cells/mL "
        f"in chemically defined, animal-component-free medium. "
    )
    if _is_numeric(titer):
        titer_val = float(titer)
        up_para += (
            f"A biphasic temperature strategy is employed with a hypothermic shift from "
            f"37°C to 33°C on culture day {d['temp_shift_day']}, which redirects cellular "
            f"energy from growth to protein production. This strategy yields a harvest titer "
            f"of {titer_val:.2f} g/L "
        )
        if titer_val >= 5.0:
            up_para += (
                f"— a commercially viable level that supports efficient downstream recovery "
                f"and favorable cost of goods. "
            )
        elif titer_val >= 2.0:
            up_para += (
                f"— an acceptable level for clinical-stage manufacturing, though further "
                f"process optimization (media enrichment, feeding strategy) may be explored "
                f"to improve commercial economics. "
            )
        else:
            up_para += (
                f"— a level that may present challenges for commercial-scale economics. "
                f"Process intensification strategies including concentrated fed-batch or "
                f"perfusion culture should be evaluated. "
            )
    if _is_numeric(vcd):
        up_para += (
            f"Peak viable cell density reaches {_num(vcd, '.1f')} x10^6 cells/mL, "
        )
    if _is_numeric(viab):
        viab_val = float(viab)
        if viab_val >= 80:
            up_para += (
                f"with harvest viability of {viab_val:.0f}%, indicating healthy culture "
                f"conditions and minimal product quality impact from cell lysis."
            )
        else:
            up_para += (
                f"with harvest viability of {viab_val:.0f}%. This relatively low viability "
                f"may increase host cell protein (HCP) and DNA burden on downstream "
                f"purification. Earlier harvest timing should be evaluated."
            )
    paras.append(up_para)

    # Scale-up context
    if d.get("su_titer") and d.get("su_strategy"):
        paras.append(
            f"Tech transfer modeling from 2L bench-scale to 2000L manufacturing-scale "
            f"bioreactor indicates a predicted titer of {d['su_titer']:.2f} g/L at "
            f"production scale using a {d['su_strategy']} scale-up strategy. This accounts "
            f"for scale-dependent effects including mixing heterogeneity, oxygen transfer "
            f"limitations, and potential shear-induced productivity losses. Engineering "
            f"parameters (P/V, kLa, tip speed) have been evaluated to ensure CHO cell "
            f"viability is maintained within the shear-safe operating window."
        )

    # Downstream narrative
    ds_para = (
        "The downstream purification process consists of a two-step chromatography "
        "platform: Protein A affinity capture followed by cation-exchange (CEX) "
        "polishing chromatography. "
    )
    if _is_numeric(ph) and _is_numeric(grad):
        ds_para += (
            f"In-silico Design of Experiments (DoE) optimization identified optimal "
            f"CEX elution conditions at pH {_num(ph, '.1f')} with a salt gradient of "
            f"{_num(grad, '.1f')} mM/min. "
        )
    if _is_numeric(res):
        res_val = float(res)
        if res_val >= 1.5:
            ds_para += (
                f"The chromatographic resolution (Rs = {res_val:.2f}) achieves baseline "
                f"separation between the main product peak and charge variants, ensuring "
                f"high purity with acceptable yield. "
            )
        elif res_val >= 0.8:
            ds_para += (
                f"The chromatographic resolution (Rs = {res_val:.2f}) provides partial "
                f"separation of charge variants. While this is adequate for clinical-stage "
                f"material, process optimization through pH fine-tuning or alternative "
                f"stationary phases (e.g., mixed-mode chromatography) may improve "
                f"separation for commercial manufacturing. "
            )
        else:
            ds_para += (
                f"The chromatographic resolution (Rs = {res_val:.2f}) indicates incomplete "
                f"separation of product-related variants. Alternative purification strategies "
                f"— including multimodal chromatography, charge engineering of the molecule, "
                f"or hydrophobic interaction chromatography (HIC) as an orthogonal step "
                f"— should be explored to achieve acceptable purity. "
            )
    if isinstance(ds_yield, (int, float)):
        ds_para += f"The overall downstream yield is {ds_yield:.1%}. "
    paras.append(ds_para)

    # COGS
    if _is_numeric(cogs):
        cogs_val = float(cogs)
        rating = d["cost_rating"]
        if cogs_val < 100:
            cogs_text = (
                f"The estimated cost of goods sold (COGS) is ${cogs_val:.2f}/g purified "
                f"drug substance ({rating}), which is highly competitive for a monoclonal "
                f"antibody product and supports favorable health-economic positioning."
            )
        elif cogs_val < 500:
            cogs_text = (
                f"The estimated COGS of ${cogs_val:.2f}/g ({rating}) is within the "
                f"typical range for commercial mAb manufacturing. Cost optimization "
                f"opportunities include titer improvement and process intensification."
            )
        else:
            cogs_text = (
                f"The current COGS estimate of ${cogs_val:.2f}/g ({rating}) is elevated "
                f"and may impact commercial viability. Key cost drivers should be identified "
                f"and addressed through process optimization."
            )
        paras.append(cogs_text)

    header = """### 3.2.S.2 Manufacture

#### 3.2.S.2.2 Description of Manufacturing Process

"""
    return header + "\n\n".join(paras)


def _narrate_characterisation(d: Dict) -> str:
    """Generate 3.2.S.3 — Characterisation narrative."""
    paras = []

    paras.append(
        f"The primary structure of {d['name']} was confirmed by the calculated intact "
        f"mass of {d['mw']} kDa, consistent with the expected molecular weight for an "
        f"IgG1 antibody of {d['seq_length']} amino acid residues including post-translational "
        f"modifications (N-linked glycosylation at Fc-N297). Identity will be further "
        f"confirmed by peptide mapping using tryptic digestion followed by LC-MS/MS, "
        f"providing >98% sequence coverage."
    )

    paras.append(
        "Higher-order structure was assessed computationally. Disulfide bond connectivity "
        "follows the canonical IgG1 pattern: 4 inter-chain disulfides (2 in the hinge region) "
        "and 12 intra-chain disulfides (4 per heavy chain, 2 per light chain). Experimental "
        "confirmation via non-reducing/reducing CE-SDS and disulfide mapping is planned."
    )

    # Impurity narrative
    imp_para = "Product-related impurities include charge variants (acidic and basic species) "
    if _is_numeric(d["agg_risk"]):
        agg_val = float(d["agg_risk"])
        if agg_val > 0.5:
            imp_para += (
                f"and high-molecular-weight (HMW) aggregates. The in-silico aggregation risk "
                f"score of {agg_val:.2f} indicates a significant propensity for HMW formation. "
                f"SEC-HPLC analysis at each process step will be employed to ensure HMW species "
                f"remain below the 5% acceptance criterion. Root cause analysis of aggregation "
                f"drivers (hydrophobic patches, thermal instability) has been performed via "
                f"the ProtePilot developability assessment module."
            )
        else:
            imp_para += (
                f"and high-molecular-weight aggregates. The predicted aggregation risk "
                f"({agg_val:.2f}) is within acceptable limits, suggesting standard process "
                f"controls and formulation conditions will maintain HMW below 5%."
            )
    else:
        imp_para += "and high-molecular-weight aggregates. Quantitative assessment is pending."
    paras.append(imp_para)

    paras.append(
        "Process-related impurities (host cell proteins, residual DNA, leached Protein A) "
        "will be controlled through validated downstream unit operations and lot release "
        "testing per ICH Q6B specifications."
    )

    header = """### 3.2.S.3 Characterisation

#### 3.2.S.3.1 Elucidation of Structure

"""
    return header + "\n\n".join(paras)


def _narrate_drug_product(d: Dict) -> str:
    """Generate 3.2.P — Drug Product narrative."""
    exc_str = ", ".join(d["excipients"]) if d["excipients"] else "None specified"
    gap = _safe_ph_pi_gap(d["pI"], d["form_ph"])
    paras = []

    # 3.2.P.1 Description
    desc_para = (
        f"The drug product is formulated as a sterile, preservative-free solution for "
        f"subcutaneous or intravenous injection."
    )
    if d["form_buffer"] != "[pending]":
        desc_para += (
            f" The formulation utilizes a {d['form_buffer']} buffer system at pH {d['form_ph']}. "
        )
    if d["excipients"]:
        exc_purposes = {
            "trehalose": "trehalose (lyoprotectant/tonicity agent)",
            "sucrose": "sucrose (stabilizer/tonicity agent)",
            "ps80": "polysorbate 80 (surfactant for aggregation prevention)",
            "polysorbate 80": "polysorbate 80 (surfactant)",
        }
        exc_detailed = [exc_purposes.get(e.lower(), e) for e in d["excipients"]]
        desc_para += f"Stabilizing excipients include {', '.join(exc_detailed)}. "
    desc_para += (
        "The primary container closure system consists of a Type I borosilicate glass "
        "vial with a fluoropolymer-coated chlorobutyl elastomeric stopper."
    )
    paras.append(desc_para)

    # 3.2.P.2 Pharmaceutical Development
    dev_para = (
        "Formulation development was guided by the ProtePilot Digital Twin platform, "
        "which integrates molecular characterization data with predictive models for "
        "colloidal stability, viscosity, and thermal degradation. "
    )
    if _is_numeric(d["pI"]) and _is_numeric(d["form_ph"]):
        gap_val = float(gap) if gap != "[pending]" else 0
        if gap_val >= 1.5:
            dev_para += (
                f"The formulation pH ({d['form_ph']}) was selected to provide a pH-pI gap of "
                f"{gap} units (pI = {d['pI']}), exceeding the minimum 1.5-unit threshold "
                f"required for adequate colloidal stability. At this gap, electrostatic "
                f"repulsion between protein molecules effectively prevents concentration-"
                f"dependent aggregation."
            )
        elif gap_val > 0:
            dev_para += (
                f"The pH-pI gap of {gap} units (formulation pH {d['form_ph']}, pI {d['pI']}) "
                f"is below the preferred 1.5-unit threshold, indicating a potential risk of "
                f"reduced colloidal stability near the isoelectric point. Additional "
                f"stabilization through increased surfactant concentration or alternative "
                f"buffer systems should be evaluated during accelerated stability studies."
            )
    paras.append(dev_para)

    # ADA/Immunogenicity integration
    if d.get("ada_risk_level"):
        ada_para = (
            f"Immunogenicity risk assessment using in-silico MHC-II peptide scanning "
            f"determined an ADA risk level of **{d['ada_risk_level']}** "
            f"(composite score: {d['ada_risk_score']:.3f}). "
        )
        if d["ada_risk_level"] == "High":
            ada_para += (
                f"A total of {d['ada_n_hotspots']} high-affinity MHC-II binding hotspots "
                f"were identified. Deimmunization engineering or T-cell epitope removal "
                f"strategies should be considered prior to clinical development. An in-vitro "
                f"PBMC T-cell proliferation assay is recommended to confirm the in-silico "
                f"findings before IND submission."
            )
        elif d["ada_risk_level"] == "Medium":
            ada_para += (
                f"{d['ada_n_hotspots']} high-risk MHC-II binding epitopes were identified. "
                f"Standard immunogenicity monitoring via a tiered anti-drug antibody (ADA) "
                f"assay (screening, confirmatory, titer) is planned for Phase I clinical "
                f"studies per FDA Guidance for Industry on Immunogenicity Testing."
            )
        else:
            ada_para += (
                "The favorable immunogenicity profile suggests standard ADA monitoring in "
                "clinical trials will be sufficient without additional deimmunization "
                "engineering."
            )
        paras.append(ada_para)

    paras.append(
        "Drug product manufacture involves ultrafiltration/diafiltration (UF/DF) concentration "
        "of the purified bulk drug substance to the target protein concentration, followed by "
        "0.22 um sterile filtration and aseptic filling under ISO Class 5 conditions. "
        "In-process controls include protein concentration (A280), pH, osmolality, "
        "sub-visible particulate counts, and bioburden prior to sterile filtration."
    )

    header = """## 3.2.P Drug Product

### 3.2.P.1 Description and Composition

"""
    section2 = "\n\n### 3.2.P.2 Pharmaceutical Development\n\n"
    section3 = "\n\n### 3.2.P.3 Manufacture\n\n"

    result = header + paras[0]
    if len(paras) > 1:
        result += section2 + "\n\n".join(paras[1:-1])
    if len(paras) > 0:
        result += section3 + paras[-1]
    return result


# ===========================================================================
# 3. Helpers
# ===========================================================================

def _is_numeric(val) -> bool:
    """Check if a value can be converted to float."""
    if val is None or val == "[pending]" or val == "[N/A]":
        return False
    try:
        float(val)
        return True
    except (ValueError, TypeError):
        return False


def _is_numeric_gt(val, threshold: float) -> bool:
    """Check if numeric and greater than threshold."""
    if not _is_numeric(val):
        return False
    return float(val) > threshold


# ===========================================================================
# 4. LLM-Enhanced Narrative (Optional)
# ===========================================================================

def _try_llm_enhance(section_text: str, section_name: str, api_key: Optional[str]) -> str:
    """
    Optionally enhance a narrative section using OpenAI API.
    Falls back to template text if API unavailable.
    """
    if not api_key:
        return section_text

    try:
        from src.llm_copilot import call_openai

        prompt = (
            f"You are a regulatory affairs expert writing a Molecular Manufacturability & Risk Assessment Report. "
            f"Improve the following {section_name} section. Make it more precise, "
            f"add relevant ICH guideline references (Q5E, Q6B, Q8, Q11), and ensure "
            f"scientific rigor. Keep the same data points. Return only the improved text.\n\n"
            f"{section_text}"
        )
        messages = [
            {"role": "system", "content": "You are a CMC regulatory writing expert."},
            {"role": "user", "content": prompt},
        ]
        enhanced = call_openai(messages, api_key, model="gpt-4o-mini", max_tokens=2048)
        if enhanced and len(enhanced) > len(section_text) * 0.5:
            return enhanced
    except Exception as e:
        log.warning(f"LLM enhancement failed for {section_name}: {e}")

    return section_text


# ===========================================================================
# 5. Main Markdown Generator (v2.0 — Deep Narrative)
# ===========================================================================

def generate_ectd_markdown(
    session_data: Dict[str, Any],
    board_result: Optional[Any] = None,
    api_key: Optional[str] = None,
) -> str:
    """
    Generate Manufacturability & Risk Assessment as Markdown with deep scientific narratives.

    Parameters
    ----------
    session_data : Streamlit session state (or dict)
    board_result : BoardMeetingResult (optional — includes CMC Board consensus)
    api_key      : OpenAI API key for optional LLM enhancement

    Returns
    -------
    str : Formatted Markdown document with paragraph-level narratives
    """
    d = _extract_filing_data(session_data)

    sections = []

    # Header
    sections.append(f"""# Molecular Manufacturability & Risk Assessment Report

## {d['name']}

**Generated by ProtePilot Platform**
**Classification:** IND / BLA Module 3 (ICH M4Q)
**Status:** DRAFT — Requires regulatory review
**Applicable Guidelines:** ICH Q5E, Q6B, Q8(R2), Q9, Q10, Q11

---
""")

    # 3.2.S.1 — Structure (deep narrative)
    s1 = _narrate_structure(d)
    s1 = _try_llm_enhance(s1, "3.2.S.1 Structure", api_key)
    sections.append(s1)
    sections.append("\n\n---\n")

    # 3.2.S.2 — Manufacture (deep narrative)
    s2 = _narrate_manufacture(d)
    s2 = _try_llm_enhance(s2, "3.2.S.2 Manufacture", api_key)
    sections.append(s2)
    sections.append("\n\n---\n")

    # 3.2.S.3 — Characterisation (deep narrative)
    s3 = _narrate_characterisation(d)
    s3 = _try_llm_enhance(s3, "3.2.S.3 Characterisation", api_key)
    sections.append(s3)
    sections.append("\n\n---\n")

    # 3.2.P — Drug Product (deep narrative)
    sp = _narrate_drug_product(d)
    sp = _try_llm_enhance(sp, "3.2.P Drug Product", api_key)
    sections.append(sp)
    sections.append("\n\n---\n")

    # CMC Board Consensus (if available)
    if board_result:
        sections.append(f"""## Appendix A: CMC Board Assessment

**Risk Level:** {board_result.risk_level}
**Risk Assessment:** {board_result.risk_assessment}
**Deliberation Time:** {board_result.wall_time_s:.2f}s

""")
        for stmt in board_result.statements:
            sections.append(f"""### {stmt.icon} {stmt.agent_name} ({stmt.agent_role})

{stmt.statement}

""")
            if stmt.risk_flags:
                sections.append("**Risk Flags:**\n")
                for f in stmt.risk_flags:
                    sections.append(f"- {f}\n")
                sections.append("\n")
        sections.append("---\n")

    # Footer
    sections.append("""
---

*This document was auto-generated by the ProtePilot Platform and constitutes
a DRAFT regulatory filing. All data must be verified by qualified regulatory affairs
personnel before submission to health authorities. Generated in compliance with ICH M4Q
(eCTD), ICH Q5E (Comparability), ICH Q6B (Specifications), ICH Q8(R2) (Pharmaceutical
Development), ICH Q9 (Quality Risk Management), and ICH Q10 (Pharmaceutical Quality System).*
""")

    return "\n".join(sections)


# ===========================================================================
# 6. Word Document Generator (v2.0 — Deep Narrative)
# ===========================================================================

def generate_ectd_docx(
    session_data: Dict[str, Any],
    board_result: Optional[Any] = None,
    api_key: Optional[str] = None,
) -> bytes:
    """
    Generate Manufacturability & Risk Assessment as a Word .docx with deep narrative paragraphs.

    Returns bytes suitable for st.download_button.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.warning("python-docx not installed; returning markdown as bytes")
        md = generate_ectd_markdown(session_data, board_result, api_key)
        return md.encode("utf-8")

    d = _extract_filing_data(session_data)
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Title Page
    title = doc.add_heading("Molecular Manufacturability & Risk Assessment Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{d['name']}")
    run.bold = True
    run.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        "Generated by ProtePilot Platform\n"
        "IND / BLA Module 3 (ICH M4Q) | DRAFT — Requires regulatory review"
    )
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # -- 3.2.S Drug Substance --
    doc.add_heading("3.2.S Drug Substance", level=1)
    doc.add_heading("3.2.S.1 General Information — Structure and Properties", level=2)

    # Structure table
    table = doc.add_table(rows=9, cols=2, style="Light Grid Accent 1")
    data_rows = [
        ("Molecular Weight", f"{d['mw']} kDa"),
        ("Isoelectric Point (pI)", str(d['pI'])),
        ("Total Amino Acids", str(d['seq_length'])),
        ("Number of Chains", str(d['num_chains'])),
        ("Surface Hydrophobicity", str(d['hydrophobicity'])),
        ("Glycoform Profile", str(d['glycoform'])),
        ("Developability Score", f"{d['dev_score']:.4f}" if isinstance(d['dev_score'], (int, float)) else str(d['dev_score'])),
        ("Aggregation Risk", f"{d['agg_risk']:.4f}" if isinstance(d['agg_risk'], (int, float)) else str(d['agg_risk'])),
        ("Thermal Stability", f"{d['stability']:.4f}" if isinstance(d['stability'], (int, float)) else str(d['stability'])),
    ]
    for i, (label, value) in enumerate(data_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Narrative paragraphs for structure
    doc.add_paragraph("")
    _add_narrative_paragraph(doc, d, "structure")

    # -- 3.2.S.2 Manufacture --
    doc.add_heading("3.2.S.2 Manufacture", level=1)
    doc.add_heading("3.2.S.2.2 Description of Manufacturing Process", level=2)

    doc.add_heading("Upstream Process (CHO Fed-Batch)", level=3)
    up_table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    up_data = [
        ("Seed Density", f"{d['seed_density']} x10^6 cells/mL"),
        ("Temperature Shift", f"Day {d['temp_shift_day']}"),
        ("Harvest Titer", f"{d['titer']} g/L" if _is_numeric(d['titer']) else str(d['titer'])),
        ("Peak VCD", f"{d['peak_vcd']} x10^6 cells/mL" if _is_numeric(d['peak_vcd']) else str(d['peak_vcd'])),
        ("Harvest Viability", f"{d['viability']}%" if _is_numeric(d['viability']) else str(d['viability'])),
    ]
    for i, (label, value) in enumerate(up_data):
        up_table.rows[i].cells[0].text = label
        up_table.rows[i].cells[1].text = value

    doc.add_paragraph("")
    _add_narrative_paragraph(doc, d, "upstream")

    ds_yield_str = f"{d['ds_yield']:.1%}" if isinstance(d['ds_yield'], (int, float)) else str(d['ds_yield'])

    doc.add_heading("Downstream Process (Protein A + IEX Polishing)", level=3)
    ds_table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    ds_data = [
        ("Optimal Elution pH", str(d['optimal_ph'])),
        ("Optimal Gradient", f"{d['optimal_gradient']} mM/min" if _is_numeric(d['optimal_gradient']) else str(d['optimal_gradient'])),
        ("Resolution (Rs)", str(d['resolution'])),
        ("Overall Yield", ds_yield_str),
        ("COGS", f"${d['cogs_per_gram']}/g ({d['cost_rating']})" if _is_numeric(d['cogs_per_gram']) else str(d['cogs_per_gram'])),
    ]
    for i, (label, value) in enumerate(ds_data):
        ds_table.rows[i].cells[0].text = label
        ds_table.rows[i].cells[1].text = value

    doc.add_paragraph("")
    _add_narrative_paragraph(doc, d, "downstream")

    # -- 3.2.S.3 Characterisation --
    doc.add_heading("3.2.S.3 Characterisation", level=1)
    doc.add_heading("3.2.S.3.1 Elucidation of Structure and Impurities", level=2)
    _add_narrative_paragraph(doc, d, "characterisation")

    # -- 3.2.P Drug Product --
    doc.add_page_break()
    doc.add_heading("3.2.P Drug Product", level=1)
    doc.add_heading("3.2.P.1 Description and Composition", level=2)

    exc_str = ", ".join(d["excipients"]) if d["excipients"] else "None specified"
    dp_table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    dp_data = [
        ("Buffer System", f"{d['form_buffer']} at pH {d['form_ph']}"),
        ("Excipients", exc_str),
        ("Dosage Form", "Solution for injection"),
        ("Container Closure", "Type I borosilicate glass vial"),
    ]
    for i, (label, value) in enumerate(dp_data):
        dp_table.rows[i].cells[0].text = label
        dp_table.rows[i].cells[1].text = value

    doc.add_paragraph("")
    _add_narrative_paragraph(doc, d, "drug_product")

    # -- ADA / Immunogenicity --
    if d.get("ada_risk_level"):
        doc.add_heading("3.2.P.2.6 Immunogenicity Risk Assessment", level=2)
        _add_narrative_paragraph(doc, d, "immunogenicity")

    # CMC Board
    if board_result:
        doc.add_page_break()
        doc.add_heading("Appendix A: CMC Board Assessment", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"Risk Level: {board_result.risk_level}")
        run.bold = True
        doc.add_paragraph(f"Risk Assessment: {board_result.risk_assessment}")

        for stmt in board_result.statements:
            doc.add_heading(f"{stmt.icon} {stmt.agent_name}", level=2)
            doc.add_paragraph(stmt.statement)
            if stmt.risk_flags:
                for f in stmt.risk_flags:
                    doc.add_paragraph(f"  - {f}")

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    run = footer.add_run(
        "This document was auto-generated by the ProtePilot Platform "
        "and constitutes a DRAFT regulatory filing. All data must be verified by "
        "qualified regulatory affairs personnel before submission to health authorities. "
        "Applicable guidelines: ICH M4Q, Q5E, Q6B, Q8(R2), Q9, Q10, Q11."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_narrative_paragraph(doc, d: Dict, section: str):
    """Add deep narrative paragraph(s) to Word document for a given section."""
    narratives = {
        "structure": _build_docx_structure_narrative,
        "upstream": _build_docx_upstream_narrative,
        "downstream": _build_docx_downstream_narrative,
        "characterisation": _build_docx_characterisation_narrative,
        "drug_product": _build_docx_drug_product_narrative,
        "immunogenicity": _build_docx_immunogenicity_narrative,
    }
    builder = narratives.get(section)
    if builder:
        for para_text in builder(d):
            doc.add_paragraph(para_text)


def _build_docx_structure_narrative(d: Dict) -> List[str]:
    paras = [
        f"{d['name']} is a recombinant monoclonal antibody (IgG) produced in CHO cells. "
        f"The molecule comprises {d['num_chains']} polypeptide chain(s), {d['seq_length']} "
        f"total residues, with a calculated intact mass of {d['mw']} kDa and theoretical "
        f"pI of {d['pI']}."
    ]
    if _is_numeric(d["dev_score"]):
        dv = float(d["dev_score"])
        risk = "favorable" if dv >= 0.65 else ("moderate" if dv >= 0.35 else "challenging")
        paras.append(
            f"The composite developability assessment yields a score of {dv:.2f}, indicating "
            f"a {risk} overall profile integrating predicted aggregation propensity, thermal "
            f"stability, and solution viscosity characteristics."
        )
    return paras


def _build_docx_upstream_narrative(d: Dict) -> List[str]:
    paras = []
    if _is_numeric(d["titer"]):
        tv = float(d["titer"])
        adequacy = "commercially competitive" if tv >= 5 else ("clinically adequate" if tv >= 2 else "requiring optimization")
        paras.append(
            f"The CHO fed-batch process employs a biphasic temperature strategy (37C to 33C "
            f"on day {d['temp_shift_day']}) to maximize specific productivity. This approach "
            f"achieves a harvest titer of {tv:.2f} g/L, which is {adequacy} for the intended "
            f"commercial manufacturing scale."
        )
    if d.get("su_titer"):
        paras.append(
            f"Scale-up modeling predicts {d['su_titer']:.2f} g/L at manufacturing scale "
            f"using {d['su_strategy']} strategy, accounting for mixing and oxygen transfer "
            f"effects."
        )
    return paras if paras else ["Upstream process data pending."]


def _build_docx_downstream_narrative(d: Dict) -> List[str]:
    paras = []
    p = "Downstream purification utilizes a Protein A capture / CEX polishing platform. "
    if _is_numeric(d["resolution"]):
        rv = float(d["resolution"])
        qual = "baseline" if rv >= 1.5 else ("partial" if rv >= 0.8 else "incomplete")
        p += f"Chromatographic resolution of Rs={rv:.2f} provides {qual} separation of charge variants."
    paras.append(p)
    return paras


def _build_docx_characterisation_narrative(d: Dict) -> List[str]:
    paras = [
        f"Primary structure is confirmed by calculated intact mass ({d['mw']} kDa). "
        f"Peptide mapping via tryptic digestion/LC-MS/MS is planned for >98% sequence coverage.",
        "Product-related impurities include charge variants and HMW aggregates controlled "
        "by CEX polishing and SEC monitoring respectively. Process-related impurities (HCP, "
        "DNA, leached Protein A) controlled through validated downstream operations."
    ]
    return paras


def _build_docx_drug_product_narrative(d: Dict) -> List[str]:
    paras = []
    exc = ", ".join(d["excipients"]) if d["excipients"] else "pending"
    gap = _safe_ph_pi_gap(d["pI"], d["form_ph"])
    paras.append(
        f"Formulation development employed the ProtePilot predictive platform. "
        f"Buffer: {d['form_buffer']} at pH {d['form_ph']} (pH-pI gap: {gap} units). "
        f"Excipients: {exc}. Drug product manufacture involves UF/DF concentration, "
        f"0.22um sterile filtration, and aseptic fill under ISO Class 5 conditions."
    )
    return paras


def _build_docx_immunogenicity_narrative(d: Dict) -> List[str]:
    if not d.get("ada_risk_level"):
        return ["Immunogenicity assessment pending."]
    paras = [
        f"In-silico immunogenicity assessment determined an ADA risk level of "
        f"{d['ada_risk_level']} (composite score: {d['ada_risk_score']:.3f}). "
        f"{d['ada_n_hotspots']} high-affinity MHC-II binding hotspots were identified. "
        f"{'Deimmunization engineering recommended. ' if d['ada_risk_level'] == 'High' else ''}"
        f"Immunogenicity monitoring via tiered ADA assay is planned per FDA guidance."
    ]
    return paras


# ===========================================================================
# 7. Self-test
# ===========================================================================

if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)
    passed = 0

    mock = {
        "last_intent": {
            "name": "TestmAb-M24",
            "pI": 8.5, "mw": 150.0,
            "sequence": "EVQLVESGGGLVQPGGSLRLSCAAS" * 18,
            "num_chains": 2, "hydrophobicity": 0.35,
        },
        "upstream_result_dict": {
            "final_titer": 9.05, "peak_vcd": 76.0,
            "viability_at_harvest": 85.0, "temp_shift_day": 5, "seed_density": 0.5,
        },
        "doe_result_dict": {
            "optimal_ph": 6.5, "optimal_gradient": 22.6,
            "optimal_resolution": 1.82, "optimal_yield": 0.82,
        },
        "cogs_result_dict": {
            "cogs_per_gram": 40.43, "cost_rating": "Excellent", "batch_output_g": 7000,
        },
        "formulation_buffer_ph": 6.0,
        "formulation_buffer_type": "histidine",
        "formulation_excipients": ["trehalose", "ps80"],
        "glycoform_profile": "standard_cho",
    }

    # Test 1: Deep narrative markdown
    md = generate_ectd_markdown(mock)
    assert "3.2.S" in md and "3.2.P" in md
    assert "TestmAb-M24" in md
    # Check for narrative prose (not just bullet points)
    assert "recombinant monoclonal antibody" in md
    assert "fed-batch" in md.lower()
    assert "colloidal stability" in md
    print(f"Test 1 PASS: Deep narrative MD = {len(md)} chars")
    passed += 1

    # Test 2: All sections present with narratives
    for section in ["3.2.S.1", "3.2.S.2", "3.2.S.3", "3.2.P.1", "3.2.P.2"]:
        assert section in md, f"Missing {section}"
    # Verify paragraph-length content (not just bullets)
    assert md.count(". ") > 20, "Expected paragraph prose, not bullet points"
    print(f"Test 2 PASS: All sections present with deep narratives ({md.count('. ')} sentences)")
    passed += 1

    # Test 3: Word document
    docx_bytes = generate_ectd_docx(mock)
    assert len(docx_bytes) > 1000
    print(f"Test 3 PASS: Word doc = {len(docx_bytes)} bytes")
    passed += 1

    # Test 4: With board result
    from src.multi_agent_board import run_cmc_board_meeting
    board = run_cmc_board_meeting(mock)
    md_board = generate_ectd_markdown(mock, board)
    assert "CMC Board" in md_board
    print(f"Test 4 PASS: MD with board = {len(md_board)} chars")
    passed += 1

    # Test 5: Empty data graceful degradation
    md_empty = generate_ectd_markdown({})
    assert "3.2.S" in md_empty
    assert len(md_empty) > 500
    print(f"Test 5 PASS: Empty data graceful = {len(md_empty)} chars")
    passed += 1

    print(f"\n{'='*50}")
    print(f"regulatory_filer v2.0 self-test: {passed}/5 passed")
    sys.exit(0 if passed == 5 else 1)
